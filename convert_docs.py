"""
Convert PDF and DOCX documents to Markdown format.
"""
import os
import sys

SOURCE_DIR = r"C:\Users\masri\OneDrive - ANSI Systems Sdn Bhd (1)\MyLaptop\01ProjectNc\LZS\NKS"
OUTPUT_DIR = r"c:\LZS - Lembaga Zakat Selangor\NKS\Docs"

FILES = [
    "Appendix G - Functional.pdf",
    "Appendix H - Technical (Application).pdf",
    "LAMPIRAN F - ARTIFICIAL INTELLIGENCE (AI).pdf",
    "LZS - Proof of Concepts (POC) V1.2.pdf",
    "Summary RFP 2026 V 1.0 .pdf",
    "NKS_2.0_AI_Proposal-4.docx",
    "NKS_Integration_3rd_Party_SRS.docx",
    "POC requirement NKS LZS.docx",
]


def pdf_to_markdown(pdf_path: str, md_path: str) -> bool:
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            lines = []
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    lines.append(text)
                    if i < len(pdf.pages) - 1:
                        lines.append("\n\n---\n\n")
            content = "".join(lines)
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(content)
            return True
    except Exception as e:
        print(f"  Error: {e}")
        return False


def docx_to_markdown(docx_path: str, md_path: str) -> bool:
    try:
        from docx import Document
        doc = Document(docx_path)
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                lines.append("")
                continue
            style = para.style.name if para.style else ""
            if "Heading 1" in style or "Title" in style:
                lines.append(f"# {text}")
            elif "Heading 2" in style:
                lines.append(f"## {text}")
            elif "Heading 3" in style:
                lines.append(f"### {text}")
            elif "Heading 4" in style:
                lines.append(f"#### {text}")
            else:
                lines.append(text)
            lines.append("")
        for table in doc.tables:
            lines.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            lines.append("")
        content = "\n".join(lines)
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(content)
        return True
    except Exception as e:
        print(f"  Error: {e}")
        return False


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    success = 0
    failed = []
    for filename in FILES:
        src = os.path.join(SOURCE_DIR, filename)
        base = os.path.splitext(filename)[0]
        md_name = base + ".md"
        dst = os.path.join(OUTPUT_DIR, md_name)
        if not os.path.exists(src):
            print(f"SKIP (not found): {filename}")
            failed.append(filename)
            continue
        print(f"Converting: {filename} -> {md_name}")
        if filename.lower().endswith(".pdf"):
            ok = pdf_to_markdown(src, dst)
        elif filename.lower().endswith(".docx"):
            ok = docx_to_markdown(src, dst)
        else:
            print(f"  Unsupported format")
            failed.append(filename)
            continue
        if ok:
            print(f"  OK: {dst}")
            success += 1
        else:
            failed.append(filename)
    print(f"\nDone: {success} converted, {len(failed)} failed")
    if failed:
        print("Failed:", ", ".join(failed))
        sys.exit(1)


if __name__ == "__main__":
    main()

